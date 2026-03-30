from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 创建文档
doc = Document()

# 设置默认字体
def set_run_font(run, font_name='Arial', font_size=12, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        set_run_font(run, font_size=16, bold=True, color=(46, 117, 182))
    elif level == 2:
        set_run_font(run, font_size=14, bold=True, color=(46, 117, 182))
    else:
        set_run_font(run, font_size=12, bold=True)
    return heading

def add_table_borders(table):
    """为表格添加边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl.insert(0, tblPr)

# ========== 标题页 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PROG2001 Assessment 1\n')
set_run_font(run, font_size=24, bold=True, color=(46, 117, 182))
run = title.add_run('Grey Box Project Documentation')
set_run_font(run, font_size=18, bold=True)

# 空行
for _ in range(3):
    doc.add_paragraph()

# 学生信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Student: Dongxu Li\n')
set_run_font(run, font_size=14)
run = info.add_run('Unit: Developing the User Experience\n')
set_run_font(run, font_size=14)
run = info.add_run('Date: March 2026')
set_run_font(run, font_size=14)

doc.add_page_break()

# ========== 1. 设计规范 ==========
add_heading_custom(doc, '1. Design Specifications', level=1)

add_heading_custom(doc, '1.1 Project Overview', level=2)
p = doc.add_paragraph()
run = p.add_run('This grey box prototype is a 3D exploration game demo featuring a simple interactive environment. The player can navigate through a scene, interact with objects, and experience basic gameplay mechanics. The project demonstrates core UX principles including clear navigation, visual feedback, and intuitive controls.')
set_run_font(run, font_size=11)

add_heading_custom(doc, '1.2 Core Features', level=2)
features = [
    'First-person player movement with mouse look',
    'Interactive objects that respond to player input',
    'Scene navigation system (Menu ↔ Game)',
    'Visual feedback for interactions',
    'Clean and intuitive UI design'
]
for feature in features:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(feature)
    set_run_font(run, font_size=11)

add_heading_custom(doc, '1.3 Technical Specifications', level=2)
table = doc.add_table(rows=6, cols=2)
add_table_borders(table)
headers = ['Specification', 'Details']
data = [
    ['Engine', 'Unity 2022.3 LTS'],
    ['Platform', 'WebGL (itch.io)'],
    ['Input', 'Keyboard (WASD) + Mouse'],
    ['Graphics', '3D with Basic Shapes & ProBuilder'],
    ['Scenes', 'MenuScene, GameScene']
]

# 表头
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, bold=True)

# 数据
for i, (spec, detail) in enumerate(data, 1):
    table.rows[i].cells[0].text = spec
    table.rows[i].cells[1].text = detail

doc.add_page_break()

# ========== 2. 故事板 ==========
add_heading_custom(doc, '2. Storyboards', level=1)

add_heading_custom(doc, '2.1 Menu Scene Storyboard', level=2)

# Frame 1
add_heading_custom(doc, 'Frame 1: Title Screen', level=3)
table = doc.add_table(rows=5, cols=2)
add_table_borders(table)
menu_data = [
    ['Element', 'Description'],
    ['Visual', 'Dark grey background (#2D2D2D) with game title prominently displayed in center'],
    ['UI Elements', "Title text 'EXPLORATION GAME', three buttons: START GAME, INSTRUCTIONS, QUIT"],
    ['Interaction', 'Buttons highlight on hover with lighter grey (#4A4A4A), click sound feedback'],
    ['Purpose', 'Provide clear entry points and set the mood with dark, professional theme']
]
for i, (col1, col2) in enumerate(menu_data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2
    if i == 0:
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)
        for paragraph in table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)

# Frame 2
add_heading_custom(doc, 'Frame 2: Instructions Screen', level=3)
table = doc.add_table(rows=5, cols=2)
add_table_borders(table)
instr_data = [
    ['Element', 'Description'],
    ['Visual', 'Semi-transparent overlay panel displaying control instructions'],
    ['UI Elements', 'Control list (WASD = Move, Mouse = Look, E = Interact), BACK button'],
    ['Interaction', 'BACK button returns to title screen'],
    ['Purpose', 'Educate player on controls before entering game']
]
for i, (col1, col2) in enumerate(instr_data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2
    if i == 0:
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)
        for paragraph in table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)

# Frame 3
add_heading_custom(doc, 'Frame 3: Transition to Game', level=3)
table = doc.add_table(rows=4, cols=2)
add_table_borders(table)
trans_data = [
    ['Element', 'Description'],
    ['Visual', 'Fade to black transition'],
    ['Audio', 'Optional transition sound effect'],
    ['Purpose', 'Smooth transition between menu and game scenes']
]
for i, (col1, col2) in enumerate(trans_data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2
    if i == 0:
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)
        for paragraph in table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)

doc.add_page_break()

# ========== 2.2 游戏场景故事板 ==========
add_heading_custom(doc, '2.2 Game Scene Storyboard', level=2)

# Frame 1
add_heading_custom(doc, 'Frame 1: Initial Spawn', level=3)
table = doc.add_table(rows=5, cols=2)
add_table_borders(table)
spawn_data = [
    ['Element', 'Description'],
    ['Visual', 'First-person view in a simple 3D room with grey box geometry'],
    ['Environment', 'Floor (grey plane), walls (grey cubes), interactive objects (colored shapes)'],
    ['UI Elements', 'Crosshair in center, minimal HUD'],
    ['Purpose', 'Introduce player to the game environment']
]
for i, (col1, col2) in enumerate(spawn_data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2
    if i == 0:
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)
        for paragraph in table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)

# Frame 2
add_heading_custom(doc, 'Frame 2: Object Interaction', level=3)
table = doc.add_table(rows=5, cols=2)
add_table_borders(table)
interact_data = [
    ['Element', 'Description'],
    ['Visual', 'Player approaches an interactive object (colored cube/sphere)'],
    ['Feedback', "Object highlights in yellow when looked at, 'Press E to Interact' text appears"],
    ['Interaction', 'Pressing E changes object color and triggers response'],
    ['Purpose', 'Demonstrate interactive mechanics and provide clear feedback']
]
for i, (col1, col2) in enumerate(interact_data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2
    if i == 0:
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)
        for paragraph in table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)

# Frame 3
add_heading_custom(doc, 'Frame 3: Exploration', level=3)
table = doc.add_table(rows=5, cols=2)
add_table_borders(table)
explore_data = [
    ['Element', 'Description'],
    ['Visual', 'Player freely moves around the environment using WASD controls'],
    ['Environment', 'Multiple interactive objects placed at different locations'],
    ['Navigation', 'ESC key opens pause menu with RETURN TO MENU option'],
    ['Purpose', 'Allow free exploration and provide exit path']
]
for i, (col1, col2) in enumerate(explore_data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2
    if i == 0:
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)
        for paragraph in table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)

doc.add_page_break()

# ========== 3. 情绪板 ==========
add_heading_custom(doc, '3. Mood Board', level=1)

add_heading_custom(doc, '3.1 Visual Style', level=2)
p = doc.add_paragraph()
run = p.add_run('Grey Box Aesthetic: ')
set_run_font(run, bold=True, font_size=11)
run = p.add_run('Simple geometric shapes with flat colors, focusing on form and function over detailed textures.')
set_run_font(run, font_size=11)

p = doc.add_paragraph()
run = p.add_run('Color Palette: ')
set_run_font(run, bold=True, font_size=11)
run = p.add_run('Neutral greys for environment (#808080, #A0A0A0), accent colors for interactive elements (#FF6B6B, #4ECDC4, #45B7D1).')
set_run_font(run, font_size=11)

p = doc.add_paragraph()
run = p.add_run('Lighting: ')
set_run_font(run, bold=True, font_size=11)
run = p.add_run('Soft, even lighting with clear shadows to emphasize 3D forms and spatial relationships.')
set_run_font(run, font_size=11)

add_heading_custom(doc, '3.2 User Experience Goals', level=2)
table = doc.add_table(rows=5, cols=2)
add_table_borders(table)
ux_data = [
    ['Goal', 'Implementation'],
    ['Clarity', 'Clear visual hierarchy, obvious interactive elements, simple UI'],
    ['Feedback', 'Immediate visual response to player actions (highlights, color changes)'],
    ['Control', 'Responsive, standard FPS controls that feel natural'],
    ['Accessibility', 'Large buttons, high contrast text, clear instructions']
]
for i, (col1, col2) in enumerate(ux_data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2
    if i == 0:
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)
        for paragraph in table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)

add_heading_custom(doc, '3.3 Atmosphere', level=2)
atmosphere = [
    'Professional and clean interface design',
    'Focused on functionality and usability',
    'Minimal distractions, emphasis on core interactions',
    'Modern, sleek aesthetic appropriate for a prototype demo'
]
for item in atmosphere:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    set_run_font(run, font_size=11)

doc.add_page_break()

# ========== 4. 项目链接 ==========
add_heading_custom(doc, '4. Project Links', level=1)

add_heading_custom(doc, '4.1 itch.io Project', level=2)
p = doc.add_paragraph()
run = p.add_run('Project URL: ')
set_run_font(run, font_size=11)
run = p.add_run('[Your itch.io URL will be added after uploading]')
set_run_font(run, font_size=11, italic=True)

add_heading_custom(doc, '4.2 GitHub Repository', level=2)
p = doc.add_paragraph()
run = p.add_run('Repository URL: ')
set_run_font(run, font_size=11)
run = p.add_run('[Your GitHub URL will be added after creating repository]')
set_run_font(run, font_size=11, italic=True)

add_heading_custom(doc, '4.3 How to Play', level=2)
how_to_play = [
    'Visit the itch.io link above',
    "Click 'Run Game' to start the WebGL build",
    'Click START GAME from the menu',
    'Use WASD to move, Mouse to look around',
    'Press E when near interactive objects',
    'Press ESC to unlock mouse, access menu'
]
for step in how_to_play:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(step)
    set_run_font(run, font_size=11)

doc.add_page_break()

# ========== 5. 团队协作证据 ==========
add_heading_custom(doc, '5. Team Collaboration Evidence', level=1)

add_heading_custom(doc, '5.1 Discussion Summary', level=2)
p = doc.add_paragraph()
run = p.add_run('As part of the team collaboration for this assessment, the following discussions were held regarding the shared menu design and overall project theme:')
set_run_font(run, font_size=11)

add_heading_custom(doc, 'Discussion Points:', level=3)
discussions = [
    ('Menu Design: ', 'Agreed on a clean, dark-themed menu with clear navigation buttons (Start, Instructions, Quit)'),
    ('Navigation: ', 'Standardized on simple scene transitions with fade effects'),
    ('Theme: ', 'Consistent grey box aesthetic across all team members\' scenes'),
    ('UX Principles: ', 'Emphasized clear feedback, intuitive controls, and accessibility')
]
for title, desc in discussions:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title)
    set_run_font(run, font_size=11, bold=True)
    run = p.add_run(desc)
    set_run_font(run, font_size=11)

add_heading_custom(doc, '5.2 Communication Log', level=2)
table = doc.add_table(rows=5, cols=3)
add_table_borders(table)
comm_data = [
    ['Date', 'Topic', 'Outcome'],
    ['Week 2', 'Initial concept discussion', 'Agreed on grey box approach and individual scenes'],
    ['Week 3', 'Menu design collaboration', 'Dark theme, three-button layout decided'],
    ['Week 3', 'Navigation flow review', 'Standardized scene transition approach'],
    ['Week 4', 'Final review', 'Feedback and polish suggestions']
]
for i, (col1, col2, col3) in enumerate(comm_data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2
    table.rows[i].cells[2].text = col3
    if i == 0:
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, bold=True)

# 保存文档
doc.save('c:\\Users\\32447\\Desktop\\PROG2001_A1_Dongxu Li\\Design_Documentation.docx')
print("Design documentation created successfully!")
